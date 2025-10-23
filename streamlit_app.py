import streamlit as st
import json
from datetime import datetime
from openai import OpenAI
import traceback
from typing import List, Dict
import tempfile
import os

# Import docx with proper error handling
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    pass  # Silently handle missing python-docx
except Exception as e:
    pass  # Silently handle docx import errors


# Show title and description.
st.title("💬 Service Advisor Review")
st.write(
    "This is a chatbot that uses OpenAI's GPT-4 model to generate responses. "
   " To use this app, upload a .mp3 voicemail file and hit 'send'."
)

# Get API key from Streamlit secrets
try:
    st.session_state.api_key = st.secrets["openai"]["api_key"]
except KeyError:
    st.error("❌ OpenAI API key not found in secrets. Please configure the openai.api_key secret.")
    st.stop()

ASSISTANT_MODEL = "gpt-4"
ASSISTANT_INSTRUCTIONS = (
    "Review uploaded voicemails between service advisors within the automotive repair industry and customers. "
    "The format should follow the instructions below and match the Mike Fix My Call Template exactly everytime, no variation or creativity beyond the inputs for each section. Use the Completed Template file in the vector store as a reference on what the final output should look like."
    "Fill out the Mike Fix My Call Template from the vector store in the assistant section with input from the audio file and reference the Sales Fix training courses "
    "when recommending next steps for improvement. Maintain a professional and positive tone and focus on creating top of the line sales teams. "
    "Reference the Example Voicemail file as a base for showing what the transcription of audio should look like and what the analysis from the coach should be. "
    "This GPT can also reference https://www.salesfix.com/ for suggestions. This GPT is only for service advisor help in the automotive industry, any other questions should result in a response of 'I am only a service advisor expert, please ask a relevant question.' "
    "When responding to the attached audio clip, automatically fill out the Mike Fix My Call Template in the vector store and format for .docx files. Remove pre and post text. The document should have bold headers for each section, the spacing should be unified with no line section breaks, the font should be unified and size 12, the title should be 'Fix My Call' with today's date and the name of the service advisor from the call. "
    "The sections for review should be 'Impression [Tonality/Charisma/Speed/Word Choice]' and 'Leadership & Professionalism [Conciseness/Confidence/Preparedness]' and 'Execution [Scripts Used/Driving Conversation/Achieved Goals]'"
    "Each section should contain notes beneath the score"
    "In the individual sections, the scoring should be highlighted red and all the options should always be displayed. The overall score should be calculated using a base of 55%, each section can add to that score for a total of 100%. "
    "A section score of Okay results in 5% increase, Good is 10% increase, and Great is 15% increase."
    "The notes area should be slightly concise, and the overall feedback slightly more verbose. Also bullet point the key point areas in the overall feedback & a summary in the feedback section. Make sure to include video names in the suggestions for follow up from the foundations advance and master series with the section titles. Also make sure that the video recommendations are listed in a sentence with comma in-between not bullet pointed."
    "The bottom should have an overall score, then state: 'Your Reviewing Trainer: Mike Tatich'"
)
ASSISTANT_TOOLS = [
    {"type": "file_search", "vector_store_ids": ["vs_68961cff51bc8191a8a5f825639a7d51"]}
]

# ---------- App state ----------
if "conversation" not in st.session_state:
    st.session_state.conversation = []
if "attached_files" not in st.session_state:
    st.session_state.attached_files = []
if "download_data" not in st.session_state:
    st.session_state.download_data = None

# ---------- Helpers ----------
def get_openai_client():
    api_key = st.session_state.get('api_key')
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


def transcribe_file(file) -> str:
    client = get_openai_client()
    if client is None:
        return "[Error: OpenAI API key not configured]"
    try:
        resp = client.audio.transcriptions.create(model="whisper-1", file=file)
        return resp.text
    except Exception as e:
        return f"[Transcription Error: {e}]"


def build_combined_user_text(user_text: str, files) -> str:
    transcripts = []
    for uploaded_file in files:
        try:
            t = transcribe_file(uploaded_file)
            transcripts.append((uploaded_file.name, t))
        except Exception as e:
            transcripts.append((uploaded_file.name, f"[TRANSCRIPTION ERROR: {e}]"))
    if transcripts:
        parts = []
        if user_text.strip():
            parts.append(user_text.strip())
        blocks = [f"--- Transcript: {fname} ---\n{text}" for fname, text in transcripts]
        parts.append("\n\n".join(blocks))
        combined = "\n\n".join(parts)
    else:
        combined = user_text.strip() or "(no text, see attachments)"
    return combined

# ---------- Responses API call ----------

def _json_schema_instruction() -> str:
    return (
        "When an audio clip is attached, output ONLY the following JSON (no prose before or after):"
        "{"
        "  \"advisor_name\": \"...\","
        "  \"date_iso\": \"YYYY-MM-DD\","
        "  \"sections\": ["
        "    {"
        "      \"name\": \"Impression [Tonality/Charisma/Speed/Word Choice]\","
        "      \"rating\": \"Needs Work|Okay|Good|Great\","
        "      \"notes\": \"coach analysis\","
        "      \"options\": [\"Needs Work\",\"Okay\",\"Good\",\"Great\"]"
        "    },"
        "    { \"name\": \"Leadership & Professionalism [Conciseness/Confidence/Preparedness]\", \"rating\": \"...\", \"notes\": \"...\", \"options\": [\"Needs Work\",\"Okay\",\"Good\",\"Great\"] },"
        "    { \"name\": \"Execution [Scripts Used/Driving Conversation/Achieved Goals]\", \"rating\": \"...\", \"notes\": \"...\", \"options\": [\"Needs Work\",\"Okay\",\"Good\",\"Great\"] },"
        "  ],"
        "  \"next_steps\": [\"...\", \"...\"],"
        "  \"transcript\": \"verbatim or cleaned transcript\""
        "}"
        "Ensure ratings reflect the call; always include all options in each section."
    )


def call_responses(conversation_messages: List[Dict], attachments_present: bool):
    client = get_openai_client()
    if client is None:
        return type("DummyResp", (), {"output_text": "[Error: OpenAI API key not configured]"})() 
    
    messages = []
    if ASSISTANT_INSTRUCTIONS:
        messages.append({"role": "system", "content": ASSISTANT_INSTRUCTIONS})
    if attachments_present:
        messages.append({"role": "system", "content": _json_schema_instruction()})
    messages.extend(conversation_messages)

    try:
        response = client.chat.completions.create(
            model=ASSISTANT_MODEL,
            messages=messages
        )
        
        # Create a response object with output_text attribute
        output_text = response.choices[0].message.content
        return type("APIResponse", (), {"output_text": output_text})()
    except Exception as e:
        return type("ErrorResp", (), {"output_text": f"[API Error: {e}]"})()

# ---------- DOCX helpers ----------
RATING_BONUS = {"Needs Work": 0, "Okay": 5, "Good": 10, "Great": 15}

def compute_overall(sections):
    total = 55
    for s in sections:
        total += RATING_BONUS.get(s.get("rating", ""), 0)
    return min(total, 100)

def _apply_body_font(run):
    run.font.size = Pt(12)

def _add_heading(p, text):
    run = p.add_run(text)
    run.bold = True
    _apply_body_font(run)

def _add_text(p, text):
    run = p.add_run(text)
    _apply_body_font(run)

def _add_red_text(p, text):
    run = p.add_run(text)
    _apply_body_font(run)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

def create_fix_my_call_docx(data: dict) -> tuple:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available")
    
    advisor = data.get("advisor_name") or "Service Advisor"
    date_iso = data.get("date_iso") or datetime.now().strftime("%Y-%m-%d")
    sections = data.get("sections", [])
    next_steps = data.get("next_steps", [])
    transcript = data.get("transcript", "")

    overall = compute_overall(sections)

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    t = title_para.add_run(f"Fix My Call — {date_iso} — {advisor}")
    t.bold = True
    t.font.size = Pt(12)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Sections
    for s in sections:
        name = s.get("name", "Section")
        rating = s.get("rating", "Needs Work")
        notes = s.get("notes", "Notes")
        opts = s.get("options") or ["Needs Work", "Okay", "Good", "Great"]

        p = doc.add_paragraph()
        _add_heading(p, name)

        p2 = doc.add_paragraph()
        for idx, opt in enumerate(opts):
            if opt == rating:
                _add_red_text(p2, opt)
            else:
                _add_text(p2, opt)
            if idx < len(opts) - 1:
                _add_text(p2, "  |  ")

        p3 = doc.add_paragraph()
        _add_text(p3, notes)

        doc.add_paragraph("")

    # Next Steps
    p = doc.add_paragraph()
    _add_heading(p, "Next Steps")
    if next_steps:
        for step in next_steps:
            li = doc.add_paragraph()
            _add_text(li, f"• {step}")
    else:
        li = doc.add_paragraph()
        _add_text(li, "• (none)")

    doc.add_paragraph("")

    # Overall Score
    p = doc.add_paragraph()
    _add_heading(p, "Overall Score: ")
    _add_red_text(p, f"{overall}%")

    doc.add_paragraph("")

    # Footer
    p = doc.add_paragraph()
    _add_heading(p, "Your Reviewing Trainer: ")
    _add_text(p, "Mike Tatich")

    safe_name = advisor.replace("/", "-").replace("\\", "-")
    filename = f"Fix My Call - {safe_name} - {date_iso}.docx"
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(temp_path)
    return temp_path, filename

# ---------- Main request ----------

def do_request(user_text: str, files):
    try:
        # Show progress bar while processing
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Debug logging
        with open("debug.log", "a", encoding="utf-8") as log:
            log.write(f"{datetime.now()}: Starting request with {len(files)} files\n")
        
        status_text.text("Processing audio file...")
        progress_bar.progress(25)
        
        combined = build_combined_user_text(user_text, files)
        st.session_state.conversation.append({"role": "user", "content": combined})

        status_text.text("Generating analysis...")
        progress_bar.progress(50)
        
        # Debug: Check API key before call
        with open("debug.log", "a", encoding="utf-8") as log:
            log.write(f"{datetime.now()}: Making API call with key: {st.session_state.api_key[:10] if st.session_state.api_key else 'None'}...\n")
        
        resp = call_responses(st.session_state.conversation, attachments_present=bool(files))
        reply = getattr(resp, "output_text", None)
        
        with open("debug.log", "a", encoding="utf-8") as log:
            log.write(f"{datetime.now()}: Got response: {reply[:100] if reply else 'None'}...\n")
            
        if not reply or "[API Error:" in reply:
            progress_bar.empty()
            status_text.text(f"Error: {reply}")
            return

        status_text.text("Creating document...")
        progress_bar.progress(75)
        
        if files and DOCX_AVAILABLE:
            try:
                with open("debug.log", "a", encoding="utf-8") as log:
                    log.write(f"{datetime.now()}: Attempting to parse JSON response\n")
                    
                data = json.loads(reply)
                if not data.get("date_iso"):
                    data["date_iso"] = datetime.now().strftime("%Y-%m-%d")
                    
                status_text.text("Finalizing document...")
                progress_bar.progress(90)
                
                with open("debug.log", "a", encoding="utf-8") as log:
                    log.write(f"{datetime.now()}: Creating DOCX with data: {data.get('advisor_name', 'Unknown')}\n")
                
                docx_path, docx_filename = create_fix_my_call_docx(data)
                
                # Complete progress
                progress_bar.progress(100)
                status_text.text("Document ready for download!")
                
                # Store download data in session state
                with open(docx_path, "rb") as f:
                    st.session_state.download_data = {
                        "data": f.read(),
                        "filename": docx_filename
                    }
                
                with open("debug.log", "a", encoding="utf-8") as log:
                    log.write(f"{datetime.now()}: Successfully created download: {docx_filename}\n")
                
                # Clear progress indicators after a moment
                import time
                time.sleep(1)
                progress_bar.empty()
                status_text.empty()
                
            except json.JSONDecodeError as je:
                with open("debug.log", "a", encoding="utf-8") as log:
                    log.write(f"{datetime.now()}: JSON parse error: {je}. Response was: {reply}\n")
                progress_bar.empty()
                status_text.text("Error: Invalid response format")
            except Exception as je:
                with open("debug.log", "a", encoding="utf-8") as log:
                    log.write(f"{datetime.now()}: DOCX creation error: {je}\n")
                progress_bar.empty()
                status_text.text(f"Error creating document: {je}")
        else:
            progress_bar.empty()
            status_text.empty()

        st.session_state.conversation.append({"role": "assistant", "content": reply})

    except Exception as e:
        tb = traceback.format_exc()
        with open("error.log", "a", encoding="utf-8") as log:
            log.write(f"{datetime.now()}: {tb}\n")
        # Clear progress indicators on error
        if 'progress_bar' in locals():
            progress_bar.empty()
        if 'status_text' in locals():
            status_text.empty()
    finally:
        st.session_state.attached_files = []

# ---------- Streamlit UI ----------
uploaded_files = st.file_uploader("Attach MP3 files", type=["mp3"], accept_multiple_files=True)
user_input = st.text_area("Type your message", height=100)

if st.button("🎯 Generate Report", type="primary"):
    if not uploaded_files:
        st.warning("⚠️ Please upload an MP3 file first.")
    else:
        # Clear any previous download state
        st.session_state.download_data = None
        do_request(user_input, uploaded_files or [])

# Show download button if data is available
if st.session_state.download_data is not None:
    st.download_button(
        label=f"📄 Download {st.session_state.download_data['filename']}",
        data=st.session_state.download_data['data'],
        file_name=st.session_state.download_data['filename'],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )

# Conversation stored but not displayed
